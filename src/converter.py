#!/usr/bin/env python
# -*- coding: utf-8 -*-
import logging
import os
import docx
import zipfile  # Potentially needed by openpyxl for error handling
from openpyxl import Workbook, load_workbook
from openpyxl.utils.exceptions import InvalidFileException
from . import utils  # 使用相对导入
from . import logger_config  # 使用相对导入

# --- Constants ---
# Word 表头（标准化后）
EXPECTED_WORD_HEADERS_NORMALIZED = [
    "序号",
    "资料名称",
    "资料来源",
    "提交人",
    "接收人",
    "交接日期",
    "存放位置",
    "备注",
]
# 期望的 Excel 输出表头 (根据用户确认的、可导入 LDIMS 的 1.xlsx 文件中的实际表头更新)
EXPECTED_EXCEL_HEADERS = [
    "文档 ID",  # 注意这里可能包含空格，以实际文件为准
    "文档名称",
    "文档类型",
    "来源部门",
    "提交人",
    "接收人",
    "签收(章)人",  # 使用实际文件中的列名
    "交接日期",  # 使用实际文件中的列名
    "保管位置",
    "备注",
    "创建人",
    "创建时间",
    "最后修改人",
    "最后修改时间",
]


class DocConverter:
    def __init__(self, word_path, excel_path):
        """
        初始化转换器。
        :param word_path: 源 Word 文档路径。
        :param excel_path: 目标 Excel 文件路径。
        """
        self.word_path = word_path
        self.excel_path = excel_path
        self.log_path = None  # 初始化为 None
        self.logger = None
        # self._setup_logger() # 将在 convert 方法开始时调用

    def _setup_logger(self):
        """配置日志记录器，日志文件与 Excel 文件同目录。"""
        try:
            # 日志文件直接放在 Excel 文件所在的目录
            log_dir = os.path.dirname(self.excel_path)
            excel_filename = os.path.splitext(os.path.basename(self.excel_path))[0]
            # 日志文件名基于 Excel 文件名
            self.log_path = os.path.join(log_dir, f"{excel_filename}_conversion.log")
            # 确保目录存在 (虽然通常dirname会存在，但以防万一)
            os.makedirs(log_dir, exist_ok=True)

            self.logger = logger_config.setup_logging(self.log_path)
            if not self.logger:
                raise RuntimeError("setup_logging returned None")
        except Exception as e:
            logging.basicConfig(
                level=logging.ERROR, format="%(asctime)s - %(levelname)s - %(message)s"
            )
            self.logger = logging.getLogger(__name__)
            # 更新错误日志中的路径信息
            intended_log_path = os.path.join(
                os.path.dirname(self.excel_path),
                f"{os.path.splitext(os.path.basename(self.excel_path))[0]}_conversion.log",
            )
            self.logger.error(
                f"Failed to setup logger at intended path '{intended_log_path}': {e}",
                exc_info=True,
            )
            self.log_path = None  # 明确设置log_path为None，表示日志设置失败

    def _check_word_table_header(self, table):
        """检查 Word 表格的表头是否符合预期（逻辑不变）。"""
        if not self.logger:
            return False  # 如果没有 logger，无法安全检查

        if not table.rows:
            self.logger.warning("Encountered a table with no rows.")
            return False
        try:
            header_cells = table.rows[0].cells
            actual_headers_normalized = [
                utils.normalize_header(cell.text) for cell in header_cells
            ]

            if len(actual_headers_normalized) != len(EXPECTED_WORD_HEADERS_NORMALIZED):
                self.logger.warning(
                    f"Table header length mismatch. Expected: {len(EXPECTED_WORD_HEADERS_NORMALIZED)}, Found: {len(actual_headers_normalized)}. Headers: {actual_headers_normalized}"
                )
                return False

            if actual_headers_normalized == EXPECTED_WORD_HEADERS_NORMALIZED:
                return True
            else:
                self.logger.warning(
                    f"Table header content mismatch. Expected: {EXPECTED_WORD_HEADERS_NORMALIZED}, Found: {actual_headers_normalized}"
                )
                return False
        except IndexError:
            self.logger.error(
                "Error accessing table header cells. The table might be malformed.",
                exc_info=True,
            )
            return False
        except Exception as e:
            self.logger.error(
                f"Unexpected error checking Word table header: {e}", exc_info=True
            )
            return False

    def _check_excel_header(self):
        """检查 Excel 文件是否存在以及表头是否匹配 (与 EXPECTED_EXCEL_HEADERS 定义比较)。"""
        if not self.logger:
            return "error"

        if not os.path.exists(self.excel_path):
            self.logger.info(
                f"Excel file '{self.excel_path}' not found. Will create a new file."
            )
            return "create"

        try:
            wb = load_workbook(self.excel_path, read_only=True)
            ws = wb.active

            if ws.max_row == 0:
                self.logger.warning(
                    f"Excel file '{self.excel_path}' exists but the active sheet is empty. Will treat as new file."
                )
                wb.close()
                return "create"

            header_row_values = next(
                ws.iter_rows(min_row=1, max_row=1, values_only=True), None
            )

            if header_row_values is None:
                self.logger.warning(
                    f"Excel file '{self.excel_path}' exists but couldn't read the header row from the active sheet. Will treat as new file."
                )
                wb.close()
                return "create"

            # 获取原始读取的表头用于日志记录和精确比较
            actual_raw_headers = [
                str(h) if h is not None else "" for h in header_row_values
            ]

            # !! 重要: 直接与代码中定义的 EXPECTED_EXCEL_HEADERS 列表进行精确比较 !!
            # 不再进行标准化或大小写转换，因为用户要求以现有文件为绝对标准

            # 比较表头长度和内容
            if len(actual_raw_headers) != len(EXPECTED_EXCEL_HEADERS):
                self.logger.error(
                    f"Excel file '{self.excel_path}' header length mismatch. "
                    f"Expected (per code): {len(EXPECTED_EXCEL_HEADERS)}, Found (in file): {len(actual_raw_headers)}. "
                    f"Expected Headers: {EXPECTED_EXCEL_HEADERS}, Found Headers: {actual_raw_headers}"
                )
                wb.close()
                return "mismatch"

            # 精确比较内容
            if actual_raw_headers == EXPECTED_EXCEL_HEADERS:
                self.logger.info(
                    f"Excel file '{self.excel_path}' exists with matching header (exact match). Will append data."
                )
                wb.close()
                return "append"
            else:
                self.logger.error(
                    f"Excel file '{self.excel_path}' header content mismatch (exact comparison). "
                    f"Expected Headers (per code): {EXPECTED_EXCEL_HEADERS}, "
                    f"Found Headers (in file): {actual_raw_headers}"
                )
                # 尝试找出第一个不匹配的位置，帮助调试
                for i, (expected, actual) in enumerate(
                    zip(EXPECTED_EXCEL_HEADERS, actual_raw_headers)
                ):
                    if expected != actual:
                        self.logger.error(
                            f"First mismatch at index {i}: Expected '{expected}', Found '{actual}'"
                        )
                        break
                wb.close()
                return "mismatch"

        except FileNotFoundError:
            self.logger.info(
                f"Excel file '{self.excel_path}' not found during header check. Will create a new file."
            )
            return "create"
        except (InvalidFileException, zipfile.BadZipFile):
            self.logger.error(
                f"Error reading Excel file '{self.excel_path}'. It might be corrupted or not a valid XLSX file.",
                exc_info=True,
            )
            return "error"
        except Exception as e:
            self.logger.error(
                f"Unexpected error reading Excel file '{self.excel_path}': {e}",
                exc_info=True,
            )
            return "error"

    def _extract_data_from_table(self, table, table_index):
        """从 Word 表格提取数据，跳过空行，并统计跳过的空行数。"""
        if not self.logger:
            return [], [], 0  # 返回空列表和计数0

        extracted_rows = []
        original_row_indices = []
        skipped_empty_count = 0  # 初始化跳过的空行计数器
        try:
            for i, row in enumerate(table.rows):
                if i == 0:  # 跳过表头行
                    continue
                # 提取原始文本
                row_data_texts = [cell.text for cell in row.cells]

                # **关键：检查是否为空行** (所有单元格文本去除空格后都为空)
                if all(not cell_text.strip() for cell_text in row_data_texts):
                    self.logger.info(
                        f"Skipping empty row found in Word table {table_index + 1}, original row index {i + 1}."
                    )
                    skipped_empty_count += 1  # 增加计数
                    continue  # 跳过空行

                # 如果不是空行，添加到结果列表
                extracted_rows.append(row_data_texts)
                original_row_indices.append(i + 1)

            # 返回提取的数据、原始行号和跳过的空行数
            return extracted_rows, original_row_indices, skipped_empty_count
        except Exception as e:
            self.logger.error(
                f"Error extracting data from table {table_index + 1}: {e}",
                exc_info=True,
            )
            return [], [], skipped_empty_count  # 出错时也返回当前计数

    def _process_row(self, raw_row_data, table_index, row_index):
        """处理单行 Word 数据，准备写入 Excel (映射到最新的14列中文表头)。"""
        if not self.logger:
            return None

        expected_word_cols = len(EXPECTED_WORD_HEADERS_NORMALIZED)
        # 检查列数是否符合 Word 表头预期 (逻辑不变)
        if len(raw_row_data) < expected_word_cols:
            self.logger.warning(
                f"Row {row_index} in table {table_index + 1} has fewer cells ({len(raw_row_data)}) than expected ({expected_word_cols}). Padding with empty strings. Data: {raw_row_data}"
            )
            raw_row_data.extend([""] * (expected_word_cols - len(raw_row_data)))
        elif len(raw_row_data) > expected_word_cols:
            self.logger.warning(
                f"Row {row_index} in table {table_index + 1} has more cells ({len(raw_row_data)}) than expected ({expected_word_cols}). Truncating extra cells. Data: {raw_row_data}"
            )
            raw_row_data = raw_row_data[:expected_word_cols]

        # 按 Word 表头顺序提取数据
        try:
            # Word 列索引: 0:序号, 1:资料名称, 2:资料来源, 3:提交人, 4:接收人, 5:交接日期, 6:存放位置, 7:备注
            data_name = raw_row_data[1].strip()  # -> 文档名称 (Excel Index 1)
            data_source = raw_row_data[2].strip()  # -> 来源部门 (Excel Index 3)
            submitter = raw_row_data[3].strip()  # -> 提交人   (Excel Index 4)
            receiver = raw_row_data[4].strip()  # -> 接收人   (Excel Index 5)
            handover_date_str = raw_row_data[5].strip()  # -> 交接日期 (Excel Index 7)
            location = raw_row_data[6].strip()  # -> 保管位置 (Excel Index 8)
            remarks = raw_row_data[7].strip()  # -> 备注     (Excel Index 9)

            # 处理日期
            handover_date_obj = utils.parse_date(handover_date_str)
            if handover_date_obj:
                handover_date_formatted = handover_date_obj.strftime("%Y-%m-%d")
            else:
                self.logger.warning(
                    f"Could not parse date '{handover_date_str}' (交接日期) in table {table_index + 1}, row {row_index}. Leaving date field empty."
                )
                handover_date_formatted = ""  # 如果日期解析失败，保留为空

            # 构建 Excel 行数据列表 (映射到14列 EXPECTED_EXCEL_HEADERS 顺序)
            excel_row = [
                "",  # 0: 文档 ID
                data_name,  # 1: 文档名称
                "",  # 2: 文档类型
                data_source,  # 3: 来源部门
                submitter,  # 4: 提交人
                receiver,  # 5: 接收人
                "",  # 6: 签收(章)人
                handover_date_formatted,  # 7: 交接日期
                location,  # 8: 保管位置
                remarks,  # 9: 备注
                "",  # 10: 创建人
                "",  # 11: 创建时间
                "",  # 12: 最后修改人
                "",  # 13: 最后修改时间
            ]

            # 验证长度 (确保内部逻辑正确)
            if len(excel_row) != len(EXPECTED_EXCEL_HEADERS):
                self.logger.error(
                    f"Internal error during row processing: Mismatched length between processed data ({len(excel_row)}) and Excel headers ({len(EXPECTED_EXCEL_HEADERS)}). Data: {excel_row}"
                )
                return None

            return excel_row

        except IndexError as e:
            self.logger.error(
                f"Index error processing row {row_index} in table {table_index + 1}. Data: {raw_row_data}. Error: {e}",
                exc_info=True,
            )
            return None
        except Exception as e:
            self.logger.error(
                f"Unexpected error processing row {row_index} in table {table_index + 1}. Data: {raw_row_data}. Error: {e}",
                exc_info=True,
            )
            return None

    def convert(self):
        """执行 Word 到 Excel 的转换过程。"""
        self._setup_logger()
        if not self.logger:
            # 即使没有文件日志，也应该能在控制台看到错误
            print("ERROR: Logger setup failed critically. Cannot proceed.")
            # 返回错误信息，避免程序完全崩溃
            return {
                "status": "error",
                "message": "Logger setup failed. Cannot proceed.",
                "success": 0,
                "errors": 0,
                "total_skipped_rows": 0,
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }

        self.logger.info(
            f"Starting conversion from '{self.word_path}' to '{self.excel_path}'"
        )

        excel_mode = self._check_excel_header()  # 调用新的检查函数

        if excel_mode == "mismatch" or excel_mode == "error":
            msg = f"Excel header check failed (mode: {excel_mode}). Please check the Excel file or logs."
            self.logger.error(msg)
            return {
                "status": "error",
                "message": msg,
                "success": 0,
                "errors": 0,
                "total_skipped_rows": 0,
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }

        success_count = 0
        error_count = 0
        total_skipped_empty = 0  # Word 读取时跳过
        processed_data_for_excel = []
        processed_tables = 0
        processed_rows_total = 0
        skipped_processed_empty_count = 0  # 处理后变空跳过

        try:
            document = docx.Document(self.word_path)
            self.logger.info(f"Successfully opened Word document: '{self.word_path}'")
            self.logger.info(f"Found {len(document.tables)} tables in the document.")

            for table_index, table in enumerate(document.tables):
                self.logger.info(f"Processing table {table_index + 1}...")
                if self._check_word_table_header(table):
                    self.logger.info(
                        f"Table {table_index + 1} header matches. Extracting data..."
                    )
                    processed_tables += 1
                    extracted_rows, original_indices, skipped_in_table = (
                        self._extract_data_from_table(table, table_index)
                    )
                    total_skipped_empty += skipped_in_table  # 累加到总数
                    self.logger.info(
                        f"Extracted {len(extracted_rows)} non-empty rows from table {table_index + 1}. Skipped {skipped_in_table} empty rows in this table."
                    )

                    if not extracted_rows:
                        continue

                    for i, raw_row in enumerate(extracted_rows):
                        original_row_index = original_indices[i]
                        processed_rows_total += 1
                        processed_row_data = self._process_row(
                            raw_row, table_index, original_row_index
                        )

                        if processed_row_data:
                            # --- 新增检查: 检查处理后的行是否有效空行 ---
                            is_processed_row_empty = all(
                                str(cell).strip() == "" for cell in processed_row_data
                            )
                            if is_processed_row_empty:
                                self.logger.warning(
                                    f"Skipping effectively empty row after processing: table {table_index + 1}, original row {original_row_index}. Raw data: {raw_row}"
                                )
                                skipped_processed_empty_count += 1  # 计数
                                # 不将此行添加到 processed_data_for_excel，也不计入 success_count 或 error_count
                            else:
                                # --- 只有非空行才添加到最终列表并计数 ---
                                success_count += 1
                                processed_data_for_excel.append(processed_row_data)
                                self.logger.debug(
                                    f"Successfully processed row: table {table_index + 1}, original row {original_row_index}."
                                )
                        else:  # _process_row 返回了 None (处理失败)
                            error_count += 1
                else:
                    self.logger.warning(
                        f"Skipping table {table_index + 1} due to header mismatch."
                    )

        except docx.opc.exceptions.PackageNotFoundError as e:
            msg = f"Word 文档未找到或无效: '{self.word_path}'"
            self.logger.error(msg)
            total_skipped_rows = (
                total_skipped_empty + skipped_processed_empty_count
            )  # 计算总数
            return {
                "status": "error",
                "message": msg,
                "success": 0,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }
        except Exception as e:
            msg = f"读取 Word 文档时发生意外错误: {e}"
            self.logger.error(msg, exc_info=True)
            total_skipped_rows = (
                total_skipped_empty + skipped_processed_empty_count
            )  # 计算总数
            return {
                "status": "error",
                "message": msg,
                "success": 0,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }

        # --- 计算总的跳过行数 ---
        total_skipped_rows = total_skipped_empty + skipped_processed_empty_count

        # --- 处理没有数据写入的情况 ---
        if not processed_data_for_excel:
            msg = "..."
            # ... (Update messages to reflect the single total_skipped_rows)
            if processed_tables == 0:
                msg = "未在 Word 文档中找到表头匹配的表格。未写入数据。"
            elif error_count > 0:
                msg = f"从 {processed_tables} 个匹配表格中处理了 {processed_rows_total} 个非空行，但 {error_count} 行处理失败，{skipped_processed_empty_count} 行处理后变为空。总共跳过 {total_skipped_rows} 行。未写入数据。请检查日志。"
            elif (
                skipped_processed_empty_count > 0
                and processed_rows_total == skipped_processed_empty_count
            ):
                msg = f"从 {processed_tables} 个匹配表格中处理了 {processed_rows_total} 个非空行，但所有行处理后均变为空。总共跳过 {total_skipped_rows} 行。未写入数据。"
            elif total_skipped_rows > 0 and processed_rows_total == 0:
                msg = f"在 {processed_tables} 个匹配表格中只找到空行 (总共跳过 {total_skipped_rows} 行)。未写入数据。"
            else:
                msg = "未从 Word 文档成功提取或处理任何数据。未写入数据。"

            self.logger.warning(msg)
            status = (
                "warning"
                if error_count == 0 and skipped_processed_empty_count == 0
                else "error"
            )
            return {
                "status": status,
                "message": msg,
                "success": 0,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }

        # --- 写入 Excel 文件 ---
        try:
            os.makedirs(os.path.dirname(self.excel_path), exist_ok=True)

            if excel_mode == "create":
                self.logger.info(f"Creating new Excel file: '{self.excel_path}'")
                wb = Workbook()
                ws = wb.active
                ws.append(EXPECTED_EXCEL_HEADERS)
                self.logger.info(
                    f"Writing header to new Excel file: {EXPECTED_EXCEL_HEADERS}"
                )
                for row_data in processed_data_for_excel:
                    ws.append(row_data)
                wb.save(self.excel_path)
                self.logger.info(
                    f"Successfully wrote {success_count} rows to new Excel file."
                )

            elif excel_mode == "append":
                self.logger.info(
                    f"Appending data to existing Excel file: '{self.excel_path}'"
                )
                wb = load_workbook(self.excel_path)
                ws = wb.active
                for row_data in processed_data_for_excel:
                    ws.append(row_data)
                wb.save(self.excel_path)
                self.logger.info(
                    f"Successfully appended {success_count} rows to Excel file."
                )

            # 简化最终消息
            final_message = f"转换完成: 成功 {success_count} 行, 失败 {error_count} 行, 共跳过空行 {total_skipped_rows} 行."
            self.logger.info(final_message)
            return {
                "status": "success",
                "message": final_message,
                "success": success_count,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }

        except PermissionError as e:
            msg = f"写入 Excel 文件 '{self.excel_path}' 失败。权限不足或文件被占用?"
            self.logger.error(msg, exc_info=False)
            return {
                "status": "error",
                "message": msg,
                "success": 0,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }
        except Exception as e:
            msg = f"写入 Excel 文件 '{self.excel_path}' 时发生意外错误: {e}"
            self.logger.error(msg, exc_info=True)
            return {
                "status": "error",
                "message": msg,
                "success": 0,
                "errors": error_count,
                "total_skipped_rows": total_skipped_rows,  # 返回总数
                "excel_path": self.excel_path,
                "log_path": self.log_path,
            }


# --- 测试块 (需要 openpyxl 来运行) ---
# if __name__ == '__main__':
#     # ... (Test block needs significant updates for new headers/mapping) ...
#     pass
